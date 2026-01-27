"""DOCX (Word document) processing."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Optional


@dataclass
class DOCXProcessingResult:
    """Result of DOCX processing."""

    original_path: Path
    markdown_path: Optional[Path] = None
    text_content: str = ""
    title: Optional[str] = None
    word_count: int = 0


class DOCXProcessor:
    """Process Word documents to markdown."""

    def __init__(self, output_dir: Optional[Path] = None):
        """Initialize DOCX processor.

        Args:
            output_dir: Directory for output files.
        """
        self.output_dir = output_dir or Path("data/processed/markdown")
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def process(self, docx_path: Path | str) -> DOCXProcessingResult:
        """Process a DOCX file to markdown.

        Args:
            docx_path: Path to DOCX file.

        Returns:
            DOCXProcessingResult with extracted content.
        """
        docx_path = Path(docx_path)

        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")

        result = DOCXProcessingResult(original_path=docx_path)

        try:
            from docx import Document

            doc = Document(docx_path)

            # Extract title from core properties
            if doc.core_properties.title:
                result.title = doc.core_properties.title

            # Convert to markdown
            markdown_parts = []

            for para in doc.paragraphs:
                style_name = para.style.name if para.style else ""
                text = para.text.strip()

                if not text:
                    continue

                # Handle headings
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.split()[-1])
                        level = min(level, 6)  # Cap at h6
                    except (ValueError, IndexError):
                        level = 1
                    markdown_parts.append(f"{'#' * level} {text}\n")

                elif style_name == "Title":
                    markdown_parts.append(f"# {text}\n")
                    if not result.title:
                        result.title = text

                elif style_name.startswith("List"):
                    # Handle list items
                    markdown_parts.append(f"- {text}")

                else:
                    # Regular paragraph
                    markdown_parts.append(f"{text}\n")

            # Handle tables
            for table in doc.tables:
                markdown_parts.append(self._table_to_markdown(table))

            result.text_content = "\n".join(markdown_parts)
            result.word_count = len(result.text_content.split())

            # Save markdown file
            if result.text_content:
                md_path = self.output_dir / f"{docx_path.stem}.md"
                md_path.write_text(result.text_content, encoding="utf-8")
                result.markdown_path = md_path

        except ImportError:
            raise ImportError("python-docx required. Install with: pip install python-docx")

        return result

    def _table_to_markdown(self, table) -> str:
        """Convert a docx table to markdown."""
        rows = []

        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

            # Add header separator after first row
            if i == 0:
                rows.append("| " + " | ".join(["---"] * len(cells)) + " |")

        return "\n".join(rows) + "\n"
